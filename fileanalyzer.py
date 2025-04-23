import os
from scapy.all import *
import collections
import ipaddress
import matplotlib.pyplot as plt
from datetime import datetime
from tkinter import Tk, filedialog
from docx import Document
from docx.shared import Inches
import tempfile
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def select_pcap_file():

    root = Tk()
    root.withdraw()
    root.attributes('-topmost', True)
    file_path = filedialog.askopenfilename(
        title="Select PCAP/PCAPNG file",
        filetypes=[("PCAP files", "*.pcap *.pcapng"), ("All files", "*.*")]
    )
    root.destroy()
    return file_path

def analyze_pcap(file_path):
  
    print(f"\nAnalyzing file: {os.path.basename(file_path)}")
    
    try:
        packets = rdpcap(file_path)
    except Exception as e:
        raise Exception(f"Error reading PCAP file: {str(e)}")
    

    stats = {
        'protocol_counts': collections.Counter(),
        'source_ips': collections.Counter(),
        'destination_ips': collections.Counter(),
        'packet_sizes': [],
        'timestamps': [],
        'conversations': collections.defaultdict(int),
        'tcp_flags': collections.Counter(),
        'dns_queries': collections.Counter(),
        'ports': {'src': collections.Counter(), 'dst': collections.Counter()},
        'total_packets': len(packets),
        'start_time': None,
        'end_time': None,
        'file_name': os.path.basename(file_path),
        'file_path': file_path,
        'file_size': f"{os.path.getsize(file_path)/1024:.2f} KB"
    }


    for packet in packets:
        try:
           
            stats['packet_sizes'].append(len(packet))
            stats['timestamps'].append(packet.time)
            
            if stats['start_time'] is None or packet.time < stats['start_time']:
                stats['start_time'] = packet.time
            if stats['end_time'] is None or packet.time > stats['end_time']:
                stats['end_time'] = packet.time

         
            if packet.haslayer(Ether):
                stats['protocol_counts']['Ethernet'] += 1
            
    
            if packet.haslayer(IP):
                stats['protocol_counts']['IP'] += 1
                src_ip = packet[IP].src
                dst_ip = packet[IP].dst
                stats['source_ips'][src_ip] += 1
                stats['destination_ips'][dst_ip] += 1
                stats['conversations'][(src_ip, dst_ip)] += 1
                
          
                if packet.haslayer(TCP):
                    stats['protocol_counts']['TCP'] += 1
                    stats['tcp_flags'][str(packet[TCP].flags)] += 1
                    stats['ports']['dst'][packet[TCP].dport] += 1
                    stats['ports']['src'][packet[TCP].sport] += 1
                    
                   
                    if packet[TCP].dport == 80 or packet[TCP].sport == 80:
                        stats['protocol_counts']['HTTP'] += 1
                    elif packet[TCP].dport == 443 or packet[TCP].sport == 443:
                        stats['protocol_counts']['HTTPS'] += 1
                        
                elif packet.haslayer(UDP):
                    stats['protocol_counts']['UDP'] += 1
                    stats['ports']['dst'][packet[UDP].dport] += 1
                    stats['ports']['src'][packet[UDP].sport] += 1
                    
                
                    if packet.haslayer(DNS) and packet[DNS].qr == 0: 
                        if packet[DNS].qd:
                            qname = packet[DNS].qd.qname.decode('utf-8', errors='ignore')
                            stats['dns_queries'][qname] += 1
                
                elif packet.haslayer(ICMP):
                    stats['protocol_counts']['ICMP'] += 1
            

            if packet.haslayer(ARP):
                stats['protocol_counts']['ARP'] += 1
            
        except Exception as e:
            print(f"Error processing packet: {str(e)}")
            continue
    

    if stats['start_time'] and stats['end_time']:
        stats['time_span'] = stats['end_time'] - stats['start_time']
    else:
        stats['time_span'] = 0
    
    return stats

def display_analysis_results(stats):
    """Display analysis results in console"""
    print("\n" + "="*60)
    print("PCAP ANALYSIS RESULTS".center(60))
    print("="*60)
    
    print(f"\n{'File Name:':<20}{stats['file_name']}")
    print(f"{'File Size:':<20}{stats['file_size']}")
    print(f"{'Total Packets:':<20}{stats['total_packets']}")
    print(f"{'Capture Duration:':<20}{stats['time_span']:.2f} seconds")
    
    print("\n" + "-"*60)
    print("PROTOCOL DISTRIBUTION".center(60))
    print("-"*60)
    for proto, count in stats['protocol_counts'].most_common():
        print(f"{proto:<15}{count:>5} packets ({count/stats['total_packets']*100:>6.1f}%)")
    
    print("\n" + "-"*60)
    print("ALL SOURCE IPs".center(60))
    print("-"*60)
    for ip, count in sorted(stats['source_ips'].items()):
        print(f"{ip:<20}{count:>5} packets")
    
    print("\n" + "-"*60)
    print("ALL DESTINATION IPs".center(60))
    print("-"*60)
    for ip, count in sorted(stats['destination_ips'].items()):
        print(f"{ip:<20}{count:>5} packets")
    
    if stats['dns_queries']:
        print("\n" + "-"*60)
        print("TOP DNS QUERIES".center(60))
        print("-"*60)
        for query, count in stats['dns_queries'].most_common(5):
            print(f"  {query:<50}{count:>5}")

def generate_plots(stats):
 

    plt.figure(figsize=(16, 10))
    
    plt.subplot(2, 2, 1)
    if stats['protocol_counts']:
        protocols, counts = zip(*stats['protocol_counts'].most_common(8))
        plt.pie(counts, labels=protocols, autopct='%1.1f%%', startangle=90)
        plt.title("Protocol Distribution", pad=20)
    
    plt.subplot(2, 2, 2)
    if stats['packet_sizes']:
        plt.hist(stats['packet_sizes'], bins=50, color='green', alpha=0.7)
        plt.title("Packet Size Distribution", pad=20)
        plt.xlabel("Size (bytes)")
        plt.ylabel("Count")
    
    plt.subplot(2, 2, 3)
    if stats['ports']['dst']:
        ports, counts = zip(*stats['ports']['dst'].most_common(8))
        plt.bar([str(p) for p in ports], counts, color='orange')
        plt.title("Top Destination Ports", pad=20)
        plt.xlabel("Port Number")
        plt.ylabel("Packet Count")
    
    plt.subplot(2, 2, 4)
    if len(stats['timestamps']) > 1:
        time_series = [(t - stats['timestamps'][0]) for t in stats['timestamps']]
        plt.plot(time_series, range(len(time_series)), 'b-')
        plt.title("Packet Arrival Timeline", pad=20)
        plt.xlabel("Time (seconds)")
        plt.ylabel("Packet Number")
    
    plt.tight_layout(pad=3.0)
    plt.show()

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os

def generate_report(stats):

    temp_images = save_plots_to_temp(stats)
    doc = Document()


    doc.add_paragraph().add_run().add_break() 

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("📘 Complete Report For the Log File")
    run.bold = True
    run.font.size = Pt(24)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("\n" * 10)
    subtitle_run = subtitle.add_run(
        f"Generated by the Log Analyser on {datetime.now().strftime('%B %d, %Y')}"
    )
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(12)

    doc.add_page_break()


    doc.add_heading('1. File Information', level=1)
    file_info = [
        ("File Name", stats["file_name"]),
        ("File Path", stats["file_path"]),
        ("File Size", stats["file_size"]),
        ("Total Packets", stats["total_packets"]),
        ("Capture Duration", f"{stats['time_span']:.2f} seconds")
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Property'
    hdr_cells[1].text = 'Value'
    for item in file_info:
        row_cells = table.add_row().cells
        row_cells[0].text = item[0]
        row_cells[1].text = str(item[1])
    doc.add_page_break()

 
    doc.add_heading('2. Protocol Distribution', level=1)
    doc.add_picture(temp_images[0], width=Inches(5.5))
    doc.add_paragraph(f"Total protocols detected: {len(stats['protocol_counts'])}")
    doc.add_page_break()


    doc.add_heading('3. IP Statistics', level=1)
    doc.add_paragraph(f"Unique Source IPs: {len(stats['source_ips'])}")
    doc.add_paragraph(f"Unique Destination IPs: {len(stats['destination_ips'])}")

    doc.add_heading('3.1 Source IPs', level=2)
    src_table = doc.add_table(rows=1, cols=2)
    src_table.style = 'Table Grid'
    hdr_cells = src_table.rows[0].cells
    hdr_cells[0].text = 'IP Address'
    hdr_cells[1].text = 'Packet Count'
    for ip, count in sorted(stats['source_ips'].items()):
        row = src_table.add_row().cells
        row[0].text = ip
        row[1].text = str(count)

    doc.add_heading('3.2 Destination IPs', level=2)
    dst_table = doc.add_table(rows=1, cols=2)
    dst_table.style = 'Table Grid'
    hdr_cells = dst_table.rows[0].cells
    hdr_cells[0].text = 'IP Address'
    hdr_cells[1].text = 'Packet Count'
    for ip, count in sorted(stats['destination_ips'].items()):
        row = dst_table.add_row().cells
        row[0].text = ip
        row[1].text = str(count)
    doc.add_page_break()


    doc.add_heading('4. Port Statistics', level=1)
    doc.add_picture(temp_images[2], width=Inches(5.5))
    doc.add_paragraph(f"Unique Source Ports: {len(stats['ports']['src'])}")
    doc.add_paragraph(f"Unique Destination Ports: {len(stats['ports']['dst'])}")
    doc.add_page_break()


    doc.add_heading('5. Packet Size Analysis', level=1)
    doc.add_picture(temp_images[1], width=Inches(5.5))
    if stats['packet_sizes']:
        avg_size = sum(stats['packet_sizes']) / len(stats['packet_sizes'])
        doc.add_paragraph(f"Average Packet Size: {avg_size:.2f} bytes")
    doc.add_page_break()


    doc.add_heading('6. Traffic Timeline', level=1)
    doc.add_picture(temp_images[3], width=Inches(5.5))


    current_datetime = datetime.now().strftime("%Y-%m-%d %H_%M_%S")
    report_path = os.path.join(
        os.path.dirname(stats['file_path']),
        f"PCAP_Report_{stats['file_name'].replace('.', '_')}_PROFESSIONAL_{current_datetime}.docx"
    )
    doc.save(report_path)
    print(f"File saved as {report_path}")


    for img in temp_images:
        try:
            os.unlink(img)
        except:
            pass

    return report_path



def save_plots_to_temp(stats):

    temp_files = []
    

    plt.figure(figsize=(6, 6))
    if stats['protocol_counts']:
        protocols, counts = zip(*stats['protocol_counts'].most_common(8))
        plt.pie(counts, labels=protocols, autopct='%1.1f%%', startangle=90)
        plt.title("Protocol Distribution")
    temp_proto = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
    plt.savefig(temp_proto.name, bbox_inches='tight')
    temp_files.append(temp_proto.name)
    plt.close()
    

    plt.figure(figsize=(6, 6))
    if stats['packet_sizes']:
        plt.hist(stats['packet_sizes'], bins=50, color='green', alpha=0.7)
        plt.title("Packet Size Distribution")
        plt.xlabel("Size (bytes)")
        plt.ylabel("Count")
    temp_size = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
    plt.savefig(temp_size.name, bbox_inches='tight')
    temp_files.append(temp_size.name)
    plt.close()

    plt.figure(figsize=(6, 6))
    if stats['ports']['dst']:
        ports, counts = zip(*stats['ports']['dst'].most_common(8))
        plt.bar([str(p) for p in ports], counts, color='orange')
        plt.title("Top Destination Ports")
        plt.xlabel("Port Number")
        plt.ylabel("Packet Count")
    temp_ports = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
    plt.savefig(temp_ports.name, bbox_inches='tight')
    temp_files.append(temp_ports.name)
    plt.close()
    

    plt.figure(figsize=(6, 6))
    if len(stats['timestamps']) > 1:
        time_series = [(t - stats['timestamps'][0]) for t in stats['timestamps']]
        plt.plot(time_series, range(len(time_series)), 'b-')
        plt.title("Packet Arrival Timeline")
        plt.xlabel("Time (seconds)")
        plt.ylabel("Packet Number")
    temp_time = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
    plt.savefig(temp_time.name, bbox_inches='tight')
    temp_files.append(temp_time.name)
    plt.close()
    
    return temp_files